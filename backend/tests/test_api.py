from __future__ import annotations

from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.analyzer.rewriter import _SuggestionList
from app.models.resume import (
    Bullet,
    Contact,
    GapReport,
    Job,
    ResumeDoc,
    Suggestion,
)
from tests.conftest import FakeAnthropic


def _docx_upload_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Sample Candidate — candidate@example.com")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Senior Engineer at Acme — Jan 2022 to Present")
    doc.add_paragraph("Was responsible for building backend services.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _seed_responses(fake: FakeAnthropic) -> None:
    bullet = Bullet(
        id="b1",
        text="Was responsible for building backend services.",
        section="experience",
        parent_id="job-1",
    )
    fake.queue(
        ResumeDoc(
            contact=Contact(name="Sample Candidate", email="candidate@example.com"),
            experience=[
                Job(
                    id="job-1",
                    company="Acme",
                    title="Senior Engineer",
                    start_date="Jan 2022",
                    end_date="Present",
                    bullets=[bullet],
                )
            ],
            raw_text="Was responsible for building backend services.",
        )
    )
    fake.queue(GapReport(ats_score=68, issues=[]))
    fake.queue(
        _SuggestionList(
            suggestions=[
                Suggestion(
                    bullet_id="b1",
                    issue="weak_verb",
                    rationale="passive voice",
                    rewrite="Built backend services that served 10K users daily.",
                )
            ]
        )
    )


def test_health_endpoint(client: TestClient) -> None:
    assert client.get("/health").json() == {"status": "ok"}


def test_upload_returns_resume_and_report(
    client: TestClient, fake_anthropic: FakeAnthropic
) -> None:
    _seed_responses(fake_anthropic)

    response = client.post(
        "/api/resume/upload",
        files={
            "file": (
                "resume.docx",
                _docx_upload_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    body = response.json()
    assert body["record_id"]
    assert body["resume"]["contact"]["name"] == "Sample Candidate"
    assert body["report"]["ats_score"] >= 0
    assert len(body["report"]["suggestions"]) == 1


def test_upload_rejects_unsupported_type(client: TestClient) -> None:
    response = client.post(
        "/api/resume/upload",
        files={"file": ("resume.txt", b"hello", "text/plain")},
    )
    assert response.status_code == 415


def test_upload_rejects_oversize(client: TestClient) -> None:
    big = b"x" * (5 * 1024 * 1024 + 1)
    response = client.post(
        "/api/resume/upload",
        files={"file": ("resume.pdf", big, "application/pdf")},
    )
    assert response.status_code == 413


def test_upload_then_export_round_trip(client: TestClient, fake_anthropic: FakeAnthropic) -> None:
    _seed_responses(fake_anthropic)
    upload_resp = client.post(
        "/api/resume/upload",
        files={
            "file": (
                "resume.docx",
                _docx_upload_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    record_id = upload_resp.json()["record_id"]
    suggestions = upload_resp.json()["report"]["suggestions"]
    suggestions[0]["accepted"] = True

    export_resp = client.post(f"/api/resume/{record_id}/export", json=suggestions)
    assert export_resp.status_code == 200
    assert (
        export_resp.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(export_resp.content) > 0

    body = Document(BytesIO(export_resp.content))
    text = "\n".join(p.text for p in body.paragraphs)
    assert "Built backend services that served 10K users daily." in text
