from io import BytesIO

from docx import Document

from app.exporter.docx_writer import build_docx
from app.models.resume import (
    Bullet,
    Contact,
    Degree,
    Job,
    Project,
    ResumeDoc,
    Suggestion,
)


def _resume() -> ResumeDoc:
    bullet = Bullet(id="b1", text="Original bullet text", section="experience", parent_id="job-1")
    return ResumeDoc(
        contact=Contact(name="Sample Candidate", email="candidate@example.com"),
        summary="Senior engineer with X years of experience.",
        experience=[
            Job(
                id="job-1",
                company="Acme",
                title="Senior Engineer",
                start_date="Jan 2022",
                end_date="Present",
                bullets=[bullet],
            )
        ],
        projects=[Project(id="p1", name="Side Project")],
        education=[Degree(id="d1", institution="IIT", qualification="BTech")],
        skills=["Python", "Go"],
    )


def _all_text(doc_bytes: bytes) -> str:
    doc = Document(BytesIO(doc_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def test_build_docx_includes_contact_and_sections() -> None:
    doc_bytes = build_docx(_resume())
    text = _all_text(doc_bytes)
    assert "Sample Candidate" in text
    assert "candidate@example.com" in text
    assert "EXPERIENCE" in text
    assert "Original bullet text" in text


def test_build_docx_applies_accepted_rewrites() -> None:
    accepted = [
        Suggestion(
            bullet_id="b1",
            issue="weak_verb",
            rationale="led",
            rewrite="Shipped X for 10K users.",
            accepted=True,
        )
    ]
    text = _all_text(build_docx(_resume(), accepted_suggestions=accepted))
    assert "Shipped X for 10K users." in text
    assert "Original bullet text" not in text


def test_build_docx_skips_unaccepted_rewrites() -> None:
    accepted = [
        Suggestion(
            bullet_id="b1",
            issue="weak_verb",
            rationale="led",
            rewrite="REWRITTEN",
            accepted=False,
        )
    ]
    text = _all_text(build_docx(_resume(), accepted_suggestions=accepted))
    assert "Original bullet text" in text
    assert "REWRITTEN" not in text


def test_build_docx_does_not_use_tables() -> None:
    """Tables are a top ATS-parsing failure mode — make sure we never emit one."""
    doc_bytes = build_docx(_resume())
    doc = Document(BytesIO(doc_bytes))
    assert len(doc.tables) == 0
