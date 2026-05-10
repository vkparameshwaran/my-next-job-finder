from __future__ import annotations

from io import BytesIO

import docx


def extract_text_from_docx(data: bytes) -> str:
    document = docx.Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
