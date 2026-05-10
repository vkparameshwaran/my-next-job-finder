"""Build a clean, ATS-friendly DOCX from a ResumeDoc.

Avoid text boxes, tables, columns, headers/footers, and unusual fonts —
those are the most common ATS parsing failures. Single-column body, standard
font (Calibri 11), bold for headings, plain bullets for lists.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.models.resume import Bullet, Degree, Job, Project, ResumeDoc, Suggestion


def _resolve_bullets(bullets: list[Bullet], accepted: dict[str, Suggestion]) -> list[str]:
    out: list[str] = []
    for b in bullets:
        suggestion = accepted.get(b.id)
        out.append(suggestion.rewrite if suggestion is not None else b.text)
    return out


def build_docx(resume: ResumeDoc, accepted_suggestions: list[Suggestion] | None = None) -> bytes:
    accepted_map: dict[str, Suggestion] = {
        s.bullet_id: s for s in (accepted_suggestions or []) if s.accepted is True
    }

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _write_header(doc, resume)
    if resume.summary:
        _write_section(doc, "Summary")
        _write_paragraph(doc, resume.summary)
    if resume.skills:
        _write_section(doc, "Skills")
        _write_paragraph(doc, ", ".join(resume.skills))
    if resume.experience:
        _write_section(doc, "Experience")
        for job in resume.experience:
            _write_job(doc, job, accepted_map)
    if resume.projects:
        _write_section(doc, "Projects")
        for project in resume.projects:
            _write_project(doc, project, accepted_map)
    if resume.education:
        _write_section(doc, "Education")
        for degree in resume.education:
            _write_degree(doc, degree, accepted_map)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _write_header(doc: Document, resume: ResumeDoc) -> None:
    contact = resume.contact
    name_para = doc.add_paragraph()
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = name_para.add_run(contact.name or "Your Name")
    run.bold = True
    run.font.size = Pt(16)

    contact_bits = [
        bit
        for bit in (
            contact.email,
            contact.phone,
            contact.location,
            contact.linkedin,
            contact.github,
        )
        if bit
    ]
    if contact_bits:
        line = doc.add_paragraph(" | ".join(contact_bits))
        line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _write_section(doc: Document, title: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)


def _write_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _write_job(doc: Document, job: Job, accepted: dict[str, Suggestion]) -> None:
    header_para = doc.add_paragraph()
    title_run = header_para.add_run(f"{job.title} — {job.company}")
    title_run.bold = True

    meta_parts = [job.location, _format_dates(job.start_date, job.end_date)]
    meta = " | ".join(p for p in meta_parts if p)
    if meta:
        meta_para = doc.add_paragraph(meta)
        meta_para.runs[0].italic = True

    for line in _resolve_bullets(job.bullets, accepted):
        doc.add_paragraph(line, style="List Bullet")


def _write_project(doc: Document, project: Project, accepted: dict[str, Suggestion]) -> None:
    header_para = doc.add_paragraph()
    name_run = header_para.add_run(project.name)
    name_run.bold = True
    if project.link:
        header_para.add_run(f" — {project.link}")
    if project.description:
        doc.add_paragraph(project.description)
    for line in _resolve_bullets(project.bullets, accepted):
        doc.add_paragraph(line, style="List Bullet")


def _write_degree(doc: Document, degree: Degree, accepted: dict[str, Suggestion]) -> None:
    header_para = doc.add_paragraph()
    qual_run = header_para.add_run(
        f"{degree.qualification}" + (f", {degree.field_of_study}" if degree.field_of_study else "")
    )
    qual_run.bold = True
    header_para.add_run(f" — {degree.institution}")

    meta_parts = [_format_dates(degree.start_date, degree.end_date), degree.grade]
    meta = " | ".join(p for p in meta_parts if p)
    if meta:
        doc.add_paragraph(meta)
    for line in _resolve_bullets(degree.bullets, accepted):
        doc.add_paragraph(line, style="List Bullet")


def _format_dates(start: str | None, end: str | None) -> str:
    if not start and not end:
        return ""
    return f"{start or ''} - {end or 'Present'}".strip(" -")
